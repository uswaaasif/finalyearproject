from transformers import BertTokenizer, BertModel,GPT2LMHeadModel, GPT2Tokenizer
from autocorrect import Speller
from docx.oxml.ns import qn
from flask import Flask, render_template, request, jsonify, send_file,redirect
from flask import send_from_directory
from werkzeug.utils import secure_filename
import os
import re
from scipy.spatial.distance import cosine
from flask import redirect, url_for
import sqlite3
from docx.text.paragraph import Paragraph
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from datetime import datetime
from docx.oxml import OxmlElement
import heapq
import nltk
from language_tool_python import LanguageTool
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from nltk.tokenize import RegexpTokenizer
from docx2pdf import convert
from nltk.corpus import wordnet
import numpy as np
import torch

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')


# Load pre-trained BERT tokenizer and model
bert_tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
bert_model = BertModel.from_pretrained('bert-base-uncased')
gpt2_tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
gpt2_model = GPT2LMHeadModel.from_pretrained("gpt2")
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def create_upload_folder():
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

def add_header_and_footer(doc, header_name):
    header = doc.sections[0].header
    header.paragraphs[0].clear()
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run(header_name).bold = True

    footer = doc.sections[0].footer
    footer.paragraphs[0].clear()
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run(datetime.now().strftime("%d %B %Y")).italic = True

def remove_extra_empty_lines(doc):
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            paragraphs_to_remove.append(paragraph)
    for paragraph in paragraphs_to_remove:
        paragraph.clear()


def generate_meaningful_title(text):
    # Tokenize the text into sentences using NLTK
    sentences = nltk.sent_tokenize(text)

    generated_titles = []

    for sentence in sentences:
        # Tokenize the sentence into words using NLTK
        words = nltk.word_tokenize(sentence)

        # Generate an initial title using BERT
        with torch.no_grad():
            inputs = bert_tokenizer(sentence, return_tensors='pt', truncation=True, padding=True)
            outputs = bert_model(**inputs)
            sentence_embedding = outputs.last_hidden_state.mean(dim=1).detach().numpy()

        # Calculate the cosine similarity between sentence embedding and each word embedding
        word_embeddings = []
        for word in words:
            with torch.no_grad():
                input_ids = torch.tensor(bert_tokenizer.encode(word, add_special_tokens=True)).unsqueeze(0)
                output = bert_model(input_ids)[0][:, 0, :].detach().numpy()
                word_embeddings.append(output)

        # Reshape sentence embedding to 1-D array
        sentence_embedding = sentence_embedding.reshape(-1)

        # Calculate cosine similarity for each word embedding
        similarities = [1 - cosine(sentence_embedding, word_embedding.flatten()) for word_embedding in word_embeddings]

        # Get the word with the highest similarity
        max_similarity_index = np.argmax(similarities)
        max_similarity_word = words[max_similarity_index]

        # Append the word to the generated title
        generated_titles.append(max_similarity_word)

    return generated_titles



def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.add_run(text)
    if style is not None:
        new_paragraph.style = style
    return new_paragraph

def add_titles_to_document(doc):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            title = generate_meaningful_title(paragraph.text)
            if title:
                title_paragraph = paragraph.insert_paragraph_before(title)
                title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in title_paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Text color black
                    run.font.underline = True
                    run.font.underline_color = RGBColor(0, 0, 255)  # Underline color blue


def format_document(doc, autocorrect=True, grammar_check=True, fix_sentences=True):
    if autocorrect:
        spell = Speller()
    if grammar_check:
        tool = LanguageTool('en-US')

    for paragraph in doc.paragraphs:
        # Remove extra spaces within the paragraph
        paragraph.text = re.sub(' +', ' ', paragraph.text)

        # Remove extra spaces at the beginning and end of the paragraph
        paragraph.text = paragraph.text.strip()

        # Correct spelling mistakes using autocorrect if enabled
        if autocorrect:
            paragraph.text = spell(paragraph.text)

        # Grammar check
        if grammar_check:
            matches = tool.check(paragraph.text)
            paragraph.text = tool.correct(paragraph.text)

        # Set paragraph font, size, and color
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'  # Change font style to Times New Roman
            run.font.size = Pt(12)  # Set font size to 12
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

        # Set paragraph alignment to justified
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Set line spacing to single
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Set indentation for all paragraphs
        paragraph.paragraph_format.left_indent = Cm(0)  # Set left indentation to 0 cm

        # Set margin from border for all paragraphs
        paragraph.paragraph_format.left_margin = Cm(1)  # Set left margin to 1 cm

        # Remove spacing before and after paragraph
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    # Generate and add titles to document
    add_titles_to_document(doc)

    # Remove extra empty lines
    remove_extra_empty_lines(doc)

    # Fix broken or incomplete sentences
    if fix_sentences:
        fix_incomplete_sentences(doc)

    # Apply border
    apply_border(doc)

def extract_summary(text):
    sentences = sent_tokenize(text)
    tokenizer = RegexpTokenizer(r'\w+')
    words = tokenizer.tokenize(text)
    stop_words = set(stopwords.words('english'))
    words = [word.lower() for word in words if word.lower() not in stop_words]
    word_freq = FreqDist(words)
    sentence_scores = {}
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in word_freq.keys():
                if len(sentence.split(' ')) < 30:  # Limiting sentence length to 30 words
                    if sentence not in sentence_scores.keys():
                        sentence_scores[sentence] = word_freq[word]
                    else:
                        sentence_scores[sentence] += word_freq[word]
    summary_sentences = heapq.nlargest(10, sentence_scores, key=sentence_scores.get)
    summary = ' '.join(summary_sentences)
    return summary

def apply_border(doc):
    width_inches = doc.sections[0].page_width / 1440
    height_inches = doc.sections[0].page_height / 1440
    left_margin_inches = doc.sections[0].left_margin.inches
    right_margin_inches = doc.sections[0].right_margin.inches
    top_margin_inches = doc.sections[0].top_margin.inches
    bottom_margin_inches = doc.sections[0].bottom_margin.inches
    content_width_inches = width_inches - left_margin_inches - right_margin_inches
    content_height_inches = height_inches - top_margin_inches - bottom_margin_inches
    is_landscape = width_inches > height_inches
    border_width = 10 if is_landscape else 20

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')

        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), str(border_width))
        top.set(qn('w:color'), '000000')
        pgBorders.append(top)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(border_width))
        bottom.set(qn('w:color'), '000000')
        pgBorders.append(bottom)

        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(border_width))
        left.set(qn('w:color'), '000000')
        pgBorders.append(left)

        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), str(border_width))
        right.set(qn('w:color'), '000000')
        pgBorders.append(right)

        sectPr.append(pgBorders)

def fix_incomplete_sentences(doc):
    for idx, paragraph in enumerate(doc.paragraphs[:-1]):
        current_paragraph = paragraph.text
        next_paragraph = doc.paragraphs[idx + 1].text
        if current_paragraph and current_paragraph[-1] not in ('.', '!', '?') and next_paragraph and next_paragraph[
            0].islower():
            doc.paragraphs[idx].add_run(' ' + next_paragraph)
            doc.paragraphs[idx + 1].clear()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload.html')
def upload_page():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    create_upload_folder()
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'})
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'})
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        header_name = request.form.get('header_name', 'Document Name')
        output_filename = 'formatted_' + filename
        doc = Document(filepath)
        add_header_and_footer(doc, header_name)
        format_document(doc)
        formatted_filepath = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        doc.save(formatted_filepath)

        # Convert to PDF
        pdf_filepath = formatted_filepath[:-4] + 'pdf'
        convert(formatted_filepath, pdf_filepath)

        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + ' '

        summary = extract_summary(text)

        return jsonify(
            {'success': 'File uploaded successfully', 'docx_filename': os.path.basename(formatted_filepath), 'summary': summary})
    else:
        return jsonify({'error': 'File type not allowed'})

@app.route('/download/<filename>')
def download_file(filename):
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(filepath):
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)
    else:
        return jsonify({'error': 'File not found'})

@app.route('/feedback.html')
def feedback_page():
    return render_template('feedback.html')
# Create a SQLite database connection
conn = sqlite3.connect('feedback.db')
cursor = conn.cursor()

# Create a table to store feedback if it doesn't exist
cursor.execute('''CREATE TABLE IF NOT EXISTS feedback (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL,
                    email TEXT NOT NULL,
                    message TEXT NOT NULL
                  )''')
conn.commit()

# Close the connection
conn.close()

# Route to handle feedback form submission
@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    try:
        # Get form data
        name = request.form['name']
        email = request.form['email']
        message = request.form['message']

        # Insert feedback into the database
        conn = sqlite3.connect('feedback.db')
        cursor = conn.cursor()
        cursor.execute('''INSERT INTO feedback (name, email, message) VALUES (?, ?, ?)''', (name, email, message))
        conn.commit()
        conn.close()

        return jsonify({'success': 'Feedback submitted successfully!'})
    except Exception as e:
        return jsonify({'error': str(e)})





if __name__ == '__main__':
    app.run(debug=True)
